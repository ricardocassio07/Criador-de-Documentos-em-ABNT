try:
    # importando um documento:
    from docx import Document
    print("✅ A biblioteca 'python-docx' está instalada! Continuando a execução...")
    # importando a função para ajustar os espaçamento entre linhas e parágrafos:
    from docx.shared import Pt
    # importando a função para ajustar o alinhamento de parágrafos:
    from docx.enum.text import WD_ALIGN_PARAGRAPH  
    # importando a função para correção do erro da fonte
    from docx.oxml.ns import qn
    # importar a função para possibilitar o recuo de todas as linhas do paragráfo
    from docx.shared import Cm
    # importar a função para possibilitar a mudança de cor da fonte:
    from docx.shared import RGBColor

    def formatacao(documento, texto, negrito=True, modelo_texto = 'nada', orientacao="centralizado", tamanho=12, pula_linhas=0, recuo_cm = 0.0, cor:int = (0,0,0)):
        """
        Adiciona um parágrafo ao documento com formatação padrão ABNT.

        Parâmetros:
            documento: objeto Document
            texto (str): texto a ser inserido
            negrito (bool): se o texto será negrito
            modelo_texto (str): 'caixaAlta' para deixar todo o texto em letras maiúsculas, 'primeiraMaiuscula' para capitalizar a primeira letra de cada palavra, ou 'nada' para manter o texto inalterado.
            orientacao (str): 'centralizado' = se o texto será centralizado, 'esquerda' = se o texto será colocado à esquerda, 'direita' = se o texto será colocado à direita, 'justificado' = será alinhado dos dois lados (justificado), 
            formando um bloco uniforme
            tamanho (int): tamanho da fonte
            pula_linhas (int): quantidade de quebras de linha antes do texto
            justificado (bool): se o texto será justicado ou não
            recuo_cm (float): recuo à esquerda em centímetros (todas as linhas do parágrafo)
            cor (int) = defina a cor do parágrafo atráves dos códigos RGB's
        """
        if modelo_texto == "caixaAlta":
            texto = texto.upper()
        elif modelo_texto == "primeiraMaiuscula":
            texto = texto.title()
        elif modelo_texto == "nada":
            texto = texto

        # Alinhamento do parágrafo
        paragrafo = documento.add_paragraph()
        if orientacao == "centralizado":
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif orientacao == "direita":
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif orientacao == "esquerda":
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif orientacao == "justificado": 
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Recuo à esquerda (todas as linha do parágrafo)
        if recuo_cm > 0:
            paragrafo.paragraph_format.left_indent = Cm(recuo_cm)

        run = paragrafo.add_run("\n" * pula_linhas + texto)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(*cor)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        run.font.size = Pt(tamanho)
        run.bold = negrito

    def verificarDesejoDoUsuario(complemento_da_pergunta:str):
        """
        Exibe uma pergunta ao usuário sobre a inclusão de um item opcional no trabalho e confirma sua decisão.

        O usuário deve responder se deseja adicionar determinado item (ex: nome da instituição, unidade ou curso).
        Em seguida, uma confirmação é solicitada. Caso o usuário confirme, a função retorna a escolha final (1 para sim, 2 para não).

        Args:
            complemento_da_pergunta (str): Texto a ser usado na pergunta inicial (ex: "o NOME DA INSTITUIÇÃO").

        Returns:
            int: Valor inteiro representando a decisão do usuário:
                - 1 se o usuário deseja adicionar o item,
                - 2 se o usuário não deseja adicionar o item.
        """
        print("\n-----------------------------------------------------------------------------------------------------------\n")
        # Variável que será responsável por parar o laço de repetição:
        loop = False
        while loop  == False:
            while True:
                try:
                    desejo = int(input(f"Você deseja adicionar {complemento_da_pergunta} no trabalho?\nDigite o número: 1- Sim 2- Não \n-> "))
                    if desejo == 1 or desejo == 2:
                        if desejo == 1:
                            decisao = "DESEJA"
                        else:
                            decisao = "NÃO DESEJA"
                        
                        verificacao = int(input(f"Tem certeza que você {decisao} adicionar {complemento_da_pergunta}?\nDigite o número: 1-SIM 2-NÃO:\n-> "))

                        if verificacao == 1:
                            loop  = True # Encerra o loop externo
                            break # Se der certo, sai do loop.
                        elif verificacao == 2:
                            print("Vamos refazer a escolha...\n")
                            break # Se der certo, sai do loop.
                        else:
                            print("Você digitou um valor inválido, lembre-se: 1-SIM 2-NÃO\n")
                    else:
                        print("Você digitou uma opção inválida!!! Lembre-se: 1-SIM 2-NÃO\n")
                except ValueError:
                    print("Você digitou um valor inválido! Por favor, digite apenas números.\n")

        if desejo == 1:
            print(f"Ok!!! Você DESEJA adicionar {complemento_da_pergunta}!!!")
        else:
            print(f"Ok!!! Você NÃO DESEJA adicionar {complemento_da_pergunta}")   

        return desejo

    def decidirMomentoDeAdicao(objeto:str):
        """
        Pergunta ao usuário se ele deseja adicionar um determinado item agora ou mais tarde, com confirmação.

        Exibe uma mensagem interativa solicitando que o usuário escolha entre adicionar o item no momento atual
        ou deixar para depois. Após a escolha inicial, a função solicita uma confirmação. Se o usuário confirmar,
        a decisão é retornada; caso contrário, a pergunta é refeita.

        Args:
            objeto (str): Nome ou descrição do item que será apresentado na pergunta (ex: "o nome do orientador").
            PRINT-> Você deseja adicionar {objeto} agora ou mais tarde?
                    Digite o número: 1- Agora 2- Mais Tarde 
                    ->

        Returns:
            int: Valor representando a decisão do usuário:
                - 1: deseja adicionar o item agora;
                - 2: deseja adicionar o item mais tarde.
        """
        print("\n")
        # Variável que será responsável por parar o laço de repetição:
        loop = False
        while loop == False:
            while True:
                try:
                    desejo = int(input(f"Você deseja adicionar {objeto} agora ou mais tarde?\nDigite o número: 1- Agora 2- Mais Tarde \n-> "))
                    if desejo == 1 or desejo == 2:
                        if desejo == 1:
                            momento = "AGORA"
                        else:
                            momento = "MAIS TARDE"
                        
                        verificacao = int(input(f"Tem certeza que você deseja adicionar {objeto} {momento}?\nDigite o número: 1-SIM 2-NÃO:\n-> "))

                        if verificacao == 1:
                            loop  = True # Encerra o loop externo
                            break # Se der certo, sai do loop.
                        elif verificacao == 2:
                            print("Vamos refazer a escolha...\n")
                            break # Se der certo, sai do loop.
                        else:
                            print("Você digitou um valor inválido, lembre-se: 1-SIM 2-NÃO\n")
                    else:
                        print("Você digitou uma opção inválida!!! Lembre-se: 1-SIM 2-NÃO\n")
                except ValueError:
                    print("Você digitou um valor inválido! Por favor, digite apenas números.\n")
        if desejo == 2:
            print(f"Ok! Você deseja adicionar {objeto} em outro momento em seu documento!\n")
    
        return desejo

    # Armazenando o arquivo com as margens personalizadas no padrão ABNT na variável global:
    documento = Document("basePythonAbnt.docx")

    # PÁGINA COM ORIENTAÇÕES PARA O USUÁRIO:
    # ESSA PÁGINA CONTERÁ DICAS E O ORIENTAÇÕES ESPECÍFICAS DO QUE O USUÁRIO DEVE MUDAR NO DOCUMENTO FINAL, O SEU TEXTO SERÁ ESCRITA EM NEGRITO E NA COR VERMELHA E NO FINAL DEVERÁ SER COMPLETAMENTE APAGADA POIS NÃO FAZ PARTE DO DOCUMENTO ABNT ORIGINAL:

    orientacoes = (
        "- Todo o conteúdo deste documento deve ser cuidadosamente revisado antes da entrega final.\n"
        "- Todos os textos destacados em vermelho devem ser corrigidos, reescritos ou apagados.\n"
        "- As seções como SUMÁRIO, LISTA DE ILUSTRAÇÕES, LISTA DE ABREVIATURAS, entre outras, devem ser feitas manualmente pelo usuário.\n"
        "- Verifique o posicionamento dos elementos nas páginas: o NOME DO LOCAL e o ANO DE ENTREGA devem estar ao final da CAPA; já o NOME DA INSTITUIÇÃO, DA UNIDADE e DO CURSO devem aparecer no topo da CAPA. Os NOMES DOS AUTORES devem ser inseridos no início da FOLHA DE ROSTO.\n"
        "- Lembre-se de criar as divisões corretamente para o cálculo do sumário. As folhas devem ser CONTADAS a partir da folha de rosto, mas NUMERADAS apenas a partir da introdução, no canto superior direito da página.\n"
        "- Use fonte Arial ou Times New Roman, tamanho 12, com espaçamento de 1,5 entre linhas em todo o texto (exceto em citações longas, notas de rodapé, referências, legendas e ficha catalográfica, que usam espaçamento simples).\n"
        "- O texto deve ser justificado, com recuo de 1,25 cm na primeira linha de cada parágrafo.\n"
        "- As margens devem seguir o padrão ABNT: 3 cm à esquerda e em cima, 2 cm à direita e embaixo.\n"
        "- Elementos obrigatórios: CAPA, FOLHA DE ROSTO, FOLHA DE APROVAÇÃO, RESUMO EM PORTUGUÊS, RESUMO EM LÍNGUA ESTRANGEIRA, SUMÁRIO e INTRODUÇÃO.\n"
        "- Todas as citações e referências devem estar de acordo com as normas da ABNT. Qualquer citação direta ou indireta deve estar corretamente referenciada ao final do trabalho.\n"
        "- Citações com mais de três linhas devem ser destacadas em parágrafo separado, com recuo de 4 cm da margem esquerda, fonte tamanho 10, sem aspas e com espaçamento simples.\n"
        "- Notas de rodapé devem ser usadas com moderação, em fonte tamanho 10 e espaçamento simples.\n"
        "- Referências devem ser listadas em ordem alfabética, com espaçamento simples entre linhas e separação de uma linha em branco entre cada item.\n"
        "- Tabelas, gráficos e imagens devem conter título numerado acima do elemento e fonte (se houver) abaixo.\n"
        "- O resumo em língua estrangeira deve ser, preferencialmente, em inglês (abstract), mas pode ser em espanhol (resumen) ou francês (résumé), conforme exigência da instituição.\n"
        "- Evite o uso da primeira pessoa do singular ('eu') ou plural ('nós') no corpo do texto. Prefira uma linguagem impessoal.\n"
        "- Atente-se à formatação dos títulos segundo os seus níveis:\n\n"
        "| Nível | Formatação                                                                         | Exemplo                                                           |\n"
        "| ----- | ---------------------------------------------------------------------------------- | ----------------------------------------------------------------- |\n"
        "| 1º    | Centralizado, negrito, caixa alta, sem ponto                                       | 1 INTRODUÇÃO                                                      |\n"
        "| 2º    | Alinhado à esquerda, negrito, caixa alta, sem ponto                                | 1.1 FUNDAMENTAÇÃO TEÓRICA                                         |\n"
        "| 3º    | Recuado 1,25 cm, negrito, só a primeira letra maiúscula, com ponto final,          | 1.1.1 Título do terceiro nível. Texto do parágrafo inicia aqui.   |\n"
        "|       | texto na mesma linha                                                               |                                                                   |\n"
        "| 4º    | Recuado 2,5 cm, negrito e itálico, só a primeira letra maiúscula, com ponto final, | 1.1.1.1 Título do quarto nível. Texto do parágrafo inicia aqui.   |\n"
        "|       | texto na mesma linha                                                               |                                                                   |\n"
        "| 5º    | Recuado 3,75 cm, itálico (sem negrito), só a primeira letra maiúscula, com ponto   | 1.1.1.1.1 Título do quinto nível. Texto do parágrafo inicia aqui. |\n"
        "|       | final, texto na mesma linha                                                        |                                                                   |\n"
    )

    formatacao(documento, "ORIENTAÇÕES", True, "caixaAlta", "centralizado", 12, 0, 0, (255, 0, 0))
    formatacao(documento,"ESTA PÁGINA CONTÉM DICAS E ORIENTAÇÕES IMPORTANTES. TODAS AS INSTRUÇÕES DEVEM SER LIDAS E EXECUTADAS PELO USUÁRIO. AO FINAL, ESSA PÁGINA DEVE SER EXCLUÍDA, POIS NÃO FAZ PARTE DO DOCUMENTO OFICIAL NO PADRÃO ABNT.",True, "caixaAlta", "esquerda", 12, 0, 0, (255, 0, 0))

    formatacao(documento, orientacoes, True, "nada", "esquerda", 12, 0, 0, (255, 0, 0))



    # Pula para a próxima página:
    documento.add_page_break()


    # CAPA:
    # Variáveis para contabilizar quantas linhas serão puladas antes de escrever o nome do autor, o título e o nome do local de entrega do documento:
    contador_de_linhas_nome_do_autor = 6
    contador_de_linhas_titulo = 8
    contador_nome_do_local = 10

    # Mensagens de boas-vindas:
    print("------------------ SEJA MUITO BEM-VINDO(A) AO SISTEMA GERADOR DE ARQUIVOS NO PADRÃO ABNT ------------------")

    # Adicionar o nome da instituição (OPCIONAL):
    # Variável criada para posterior concatenação:
    instituicao = ""
    if verificarDesejoDoUsuario("o NOME DA INSTITUIÇÃO") == 1:
        # Modifica o texto automaticamente para CAIXA ALTA:
        instituicao = input("Certo!! Digite o nome da sua instituição:\n-> ").upper()

    # Adicionar o nome da unidade (OPCIONAL):
    # Variável criada para posterior concatenação:
    unidade = ""       
    if verificarDesejoDoUsuario("o NOME DA UNIDADE") == 1:
        # Modifica o texto automaticamente para CAIXA ALTA:
        unidade = input("Certo!! Digite o nome da sua unidade:\n-> ").upper()

    # Adicionar o nome do curso (OPCIONAL):
    # Variável criada para posterior concatenação:
    curso = ""
    if verificarDesejoDoUsuario("o NOME DO SEU CURSO") == 1:
        # Todas as primeiras letras de cada palavra é modificada automaticamente paara MIÚSCULAS: 
        curso = input("Certo!! Digite o nome do seu curso:\n-> ").title()

    nome_da_instituicao_da_unidade_e_do_curso = (f"{instituicao}\n{unidade}\n{curso}")
    formatacao(documento, nome_da_instituicao_da_unidade_e_do_curso, True, "nada")


    # Nome do autor:
    # Verificar quantos autores a obra possui:
    verificarQuantidadeDeAutores = False
    # Variável para alocar a quantidade de autores da obra:
    quantidade_de_autores = 0
    print("\n-----------------------------------------------------------------------------------------------------------\n")
    while verificarQuantidadeDeAutores  == False:
        while True:
            try:
                quantidade_de_autores = int(input("Digite a quantidade de autores que a obra possui:\n-> "))
                if quantidade_de_autores > 0:
                    verificacao = int(input(f"Tem certeza que a sua obra possui {quantidade_de_autores} autor(es)? 1-SIM  2-NÃO:\n-> "))
                    if verificacao == 1:
                        print("Ok!!!")
                        verificarQuantidadeDeAutores  = True # Encerra o loop externo
                        break # Se der certo, sai do loop.
                    elif verificacao == 2:
                        quantidade_de_autores = int(input("Certo! Então, digite a quantidade correta de autores que a obra possui:\n-> "))
                        while quantidade_de_autores <= 0:
                            print("Você digitou uma opção inválida!!! Lembre-se: O NÚMERO DE AUTORES DEVE SER MAIOR DO QUE ZERO!!!")
                            quantidade_de_autores = int(input("Digite a quantidade correta de autores que a obra possui:\n-> "))
                        print("Ok!!!")
                        verificarQuantidadeDeAutores  = True # Encerra o loop externo
                        break # Se der certo, sai do loop.
                    else:
                        print("Você digitou um valor inválido, lembre-se: 1-SIM 2-NÃO")
                else:
                    print("Você digitou uma opção inválida!!! Lembre-se: O NÚMERO DE AUTORES DEVE SER MAIOR DO QUE ZERO!!!")
            except ValueError:
                print("Você digitou um valor inválido! Por favor, digite apenas números.")

    print("\n-----------------------------------------------------------------------------------------------------------\n")
    # Lista para armazenar o nome dos autores:
    autores = []
    #  Criando um parágrafo para o(s) autor(es):
    nome_do_autor = documento.add_paragraph()
    for i in range(quantidade_de_autores):
        autores.append(input(f"Qual o nome do {i+1}º autor da obra?\n-> ").upper())  # modifica o texto para CAIXA ALTA conforme funções naturais do Python.

        # Caso seja o primeiro autor que está sendo escrito no documento, será pulada a quantidade de linhas alocadas na variável contadora, caso não, não será pulada nenhuma linha.
        if i == 0:
            formatacao(documento, autores[i], True, "caixaAlta", "centralizado", 12, contador_de_linhas_nome_do_autor)
        else:
            formatacao(documento, autores[i], True, "caixaAlta", "centralizado", 12)

        # Caso o documento tenha mais de UM autor, ou seja, i sendo maior que ZERO, irá subtrair um do total da variável contadora de linhas a serem puladas para escrever o título.
        if i > 0:
            contador_de_linhas_titulo -= 1

    print("\n-----------------------------------------------------------------------------------------------------------\n")
    # Adicionar o título da obra:
    titulo = input("Digite o título do trabalho:\n(CASO HOUVER SUBTÍTULO NÃO DIGITE-O AQUI!!!)\n-> ").upper()
    # formatacao(documento, titulo,True,True)

    # Adicionar o subtítulo da obra (OPCIONAL):
    # Variável criada para posterior verificação se a variável não está vazia e, assim, o trabalho possui um subtítulo:
    subtitulo = ""
    if verificarDesejoDoUsuario("um SUBTÍTULO") == 1:

        subtitulo = input("Certo!! Digite o SUBTÍTULO do seu trabalho:\n-> ")

        titulo_e_subtitulo_da_obra_capa = (f"{titulo}: {subtitulo}") 

        formatacao(documento, titulo_e_subtitulo_da_obra_capa, True, False, "centralizado", 12,contador_de_linhas_titulo)


    print("\n-----------------------------------------------------------------------------------------------------------\n")
    # Local:
    while True:
        try:
            # Modifica o texto automaticamente para CAIXA ALTA:
            local = input("Digite o LOCAL (CIDADE) DE ENTREGA do trabalho:\n-> ").upper()
            break
        except ValueError:
            print("Você digitou um valor inválido! Por favor, digite apenas LETRAS.")


    print("\n-----------------------------------------------------------------------------------------------------------\n")
    # Ano:
    while True:
        try:
            ano = int(input("Digite o ano de entrega do trabalho:\n-> "))
            break
        except ValueError:
            print("Você digitou um valor inválido! Por favor, digite apenas números.")
    # Após a verificação, vamos converter a variável do tipo int para str para juntarmos com o local de entrega do trabalho:
    ano = str(ano)
    # Adiciona o local e o ano no documento:
    local_e_data = (f"{local}\n{ano}")
    formatacao(documento, local_e_data, True, True, "centralizado", 12, contador_nome_do_local)

    # Pula para a próxima página:
    documento.add_page_break()

    # Próxima Página: Folha de Rosto
    # Variáveis para contabilizar quantas linhas serão puladas antes de escrever o título, a nota explicativa e o nome do local de entrega do documento:
    contador_de_linhas_titulo_folha_de_rosto = 9
    contador_de_linhas_nota_explicativa_folha_de_rosto = 9
    contador_de_linhas_ano_folha_de_rosto = 7

    # Autor:
    for i in range(quantidade_de_autores):
        formatacao(documento, autores[i], True, True, "centralizado", 12)

        # Caso tenha mais de um autor, irá subtrair uma linha do total que serão puladas para escrever o título da obra.
        if i > 0:
            contador_de_linhas_titulo_folha_de_rosto -= 1


    # Título e subtítulo:

    # Verificar se a variável subtítulo não está vazia, para isso, foi utilizado a função .strip() para retirar todos os espaços do começo ao fim e verificar se ela está vazia ou contém apenas espaços:
    '''
        - strip() remove todos os espaços do começo e do fim.

        - Se o resultado ainda tiver algo (ex: "Trabalho"), será considerado True.

        - Se for só espaços ou vazio (""), será considerado False.

    '''
    if subtitulo.strip():
        titulo_completo_folha_de_rosto = (f"{titulo}:{subtitulo}")
        formatacao(documento, titulo_completo_folha_de_rosto, True, "nada", "centralizado", 12, contador_de_linhas_titulo_folha_de_rosto)
    else:
        formatacao(documento, titulo, True, False, "centralizado", 12, contador_de_linhas_titulo_folha_de_rosto)


    print("\n-----------------------------------------------------------------------------------------------------------\n")
    # Nota explicativa:
    nota_explicativa = ""
    print(
        "Digite uma nota explicativa para o seu trabalho.\n"
        "Exemplo:\n"
        "Trabalho apresentado ao curso (técnico/superior) em (NOME DO CURSO),\n"
        "da (INSTITUIÇÃO), orientado pelo(a) Prof.(a) (NOME DO PROFESSOR),\n"
        "como requisito parcial para obtenção da aprovação em (DISCIPLINA).\n"
    )
    # Verificar se o usuário deseja utilizar o modelo proposto pelo sistema:
    # Variável para alocar a resposta final do usuário:
    utilizarModelo = 0
    # Variável que será responsável por para o laço de repetição:
    verificarDesejoModeloNota = False
    while verificarDesejoModeloNota == False:
        try:
            utilizarModelo = int(input("\nVocê deseja utilizar o modelo disponibilizado pelo sistema? \nDigite o número: 1-SIM 2-NÃO\n-> "))
            if utilizarModelo == 1 or utilizarModelo == 2:
                if utilizarModelo == 1:
                    desejo = "DESEJA"
                else:
                    desejo = "NÃO DESEJA"

                verificacao = int(input(f"Tem certeza que você {desejo} utilizar o nosso modelo no seu trablho? 1-SIM 2-NÃO:\n-> "))

                if verificacao == 1:
                    verificarDesejoModeloNota == False # Para loop externo
                    break
                elif verificacao == 2:
                    print("Então, digite novamente:")
                    break
                else:
                        print("Você digitou um valor inválido, lembre-se: 1-SIM 2-NÃO")
            else:
                print("Você digitou uma opção inválida!!! Lembre-se: 1-SIM 2-NÃO")

        except ValueError:
                print("Você digitou um valor inválido! Por favor, digite apenas números.\n")   

    if utilizarModelo == 1:
        print("Perfeito!!!\nEntão nos diga:")
        while True:
            try:
                tipo_curso = int(input("O seu curso é: 1-Técnico 2-Superior?\n-> "))
                if tipo_curso in [1, 2]:
                    if tipo_curso == 1:
                        tipo_curso = "Técnico"
                    else:
                        tipo_curso = "Superior"
                    break
                else:
                    print("Você deve digitar 1 para Técnico ou 2 para Superior.")
            except ValueError:
                print("Você digitou um valor inválido! Por favor, digite apenas números.\n")
        curso = input(f"O seu curso {tipo_curso} é em:\n-> ").title()
        instituicao = input(f"O seu curso {tipo_curso} em {curso} é realizado pela instituição chamada:\n-> ").title()
        orientador = input(f"Como se chama o professor(a) orientador(a) desse trabalho?\n-> ").title()
        materia = input(f"Qual o nome da matéria na qual esse trabalho será entrege como requisito parcial de nota?\n-> ").title()
        nota_explicativa = (f"Trabalho apresentado ao curso {tipo_curso} em {curso}, da {instituicao}, orientado pelo(a) Prof.(a) {orientador}, como requisito parcial para obtenção da aprovação em {materia}.")
        formatacao(documento, nota_explicativa, False, "nada", "justificado", 11, contador_de_linhas_nota_explicativa_folha_de_rosto, 7.5, (0, 0, 0))
    elif utilizarModelo == 2:
        momentoAdicaoEpigrafe = decidirMomentoDeAdicao("a EPÍGRAFE")
        if momentoAdicaoEpigrafe == 1:
            nota_explicativa = input("Ok!!! Então digite a sua própria NOTA EXPLICATIVA:\n->")
            formatacao(documento, nota_explicativa, False, "nada", "justificado", 11, contador_de_linhas_nota_explicativa_folha_de_rosto, 7.5, (0, 0, 0))
        elif momentoAdicaoEpigrafe == 2:
            nota_explicativa = ("Aqui deve ser escrito a nota explicativa. Segundo a ABNT, ela serve para informar o tipo de trabalho, seu objetivo acadêmico, o curso, a instituição de ensino e, opcionalmente, o orientador. Ela esclarece se o trabalho é, por exemplo, um relatório, uma monografia ou um TCC, e especifica que ele é apresentado como requisito para a obtenção de um título.")
            formatacao(documento, nota_explicativa, False, "nada", "justificado", 11, contador_de_linhas_nota_explicativa_folha_de_rosto, 7.5, (255, 0, 0))



    # Local e ano de entrega:
    formatacao(documento, local_e_data, True, "nada", "centralizado", 12, contador_de_linhas_ano_folha_de_rosto)


    # Pula para a próxima página:
    documento.add_page_break()
    # Próxima Página: Dedicatória (Opcional)

    # Verificar se o usuário deseja que o seu documento tenha uma página com uma dedicatória:
    if verificarDesejoDoUsuario("UMA PÁGINA CONTENDO UMA DEDICATÓRIA") == 1:
        # Verificar em qual momento o usuário deseja adicionar o conteúdo:
        if decidirMomentoDeAdicao("a DEDICATÓRIA") == 1:
            dedicatoria = input("Ok! Então digite a sua própria DEDICATÓRIA:\n" \
            "Exemplo: Dedico aos meus pais e amigos que me auxiliaram durante o processo de construção deste trabalho\n" \
            "-> ")
        else:
            explicacaoDedicatoria = ("A dedicatória é uma parte opcional de trabalhos acadêmicos que permite ao autor expressar agradecimentos pessoais e homenagens a pessoas ou entidades que tiveram um papel importante durante a elaboração do trabalho.\n"
            "Embora não seja uma exigência formal da ABNT, a dedicatória é tradicionalmente colocada logo após a folha de rosto, antes do resumo ou sumário, e costuma ser um texto curto, direto e afetivo. Pode ser direcionada a familiares, amigos, professores, orientadores, colegas ou até a Deus, dependendo da preferência do autor.\n"
            "A função principal da dedicatória é demonstrar gratidão, reconhecimento e valorização pelo apoio, incentivo e motivação recebidos durante o processo de estudo e pesquisa. Ela tem um tom mais pessoal e emocional, diferenciando-se do agradecimento, que costuma ser mais formal e detalhado.\n"
            "*EXEMPLO ESCRITO NO CANTO DA PÁGINA, ONDE A A DEDICATÓRIA DEVE ESTAR LOCALIZADA\n"
            "*ESSA PÁGINA NÃO CONTÉM NADA ALÉM DO TEXTO, OU SEJA, NÃO POSSUI TITULOS, IMAGENS OU AFINS...\n")
            dedicatoria = ("Dedico este trabalho à minha família, pelo apoio incondicional em todos os momentos; à minha mãe, que sempre acreditou nos meus sonhos e me ensinou a perseverar; aos meus amigos, que ofereceram palavras de incentivo e companhia; ao meu orientador, pela paciência e dedicação na orientação; e a Deus, pela força e sabedoria em cada etapa dessa jornada. Sem todos vocês, este trabalho não teria sido possível.")
            formatacao(documento, explicacaoDedicatoria, True, "nada", "esquerda", 12, 0, 0, (255, 0, 0))
        formatacao(documento, dedicatoria, False, "nada", "justificado", 12, 8, 7.5 )
        # Pula para a próxima página:
        documento.add_page_break()


    # Próxima Página: Agradecimentos (Opcional):

    # Verificar se o usuário deseja que o seu documento tenha uma página com agradecimento:
    if verificarDesejoDoUsuario("UMA PÁGINA CONTENDO UM AGRADECIMENTO") == 1:
        # Verificar em qual momento o usuário deseja adicionar o conteúdo:
        if decidirMomentoDeAdicao("o AGRADECIMENTO") == 1:
            agradecimento = input("Ok! Então digite o seu próprio AGRADECIMENTO:\n" 
            "Dica: Redija um texto mostrando o quão grato você é por quem te auxiliou a chaegar até aqui.\n" 
            "Exemplo: A Deus, pela minha vida, e por me ajudar a ultrapassar todos os obstáculos encontrados ao longo do curso.\n"
            "Aos meus pais e irmãos, que me incentivaram nos momentos difíceis e compreenderam a minha ausência enquanto eu me dedicava à realização deste trabalho.\n"
            "Aos professores, pelas correções e ensinamentos que me permitiram apresentar um melhor desempenho no meu processo de formação profissional.\n"
            "-> ")
            # O título e o texto será adicionado na cor preta:
            formatacao(documento, "AGRADECIMENTOS", True, "nada", "centralizado")
            formatacao(documento, agradecimento, False, "nada", "justificado", 12)
        else:
            explicacaoAgradecimento = ("Os agradecimentos são uma seção opcional, porém bastante comum, em trabalhos acadêmicos. Eles oferecem ao autor a oportunidade de expressar, de maneira formal e respeitosa, sua gratidão às pessoas, instituições ou entidades que contribuíram de forma significativa para a realização do trabalho.\n"
            "De acordo com as normas da ABNT, os agradecimentos são incluídos após a dedicatória (caso exista) e antes do resumo. Embora o conteúdo seja mais pessoal do que o restante do trabalho, o tom costuma ser mais sóbrio e estruturado do que o da dedicatória, mantendo uma linguagem formal e respeitosa.\n"
            "Nesse espaço, é comum agradecer a professores e orientadores pela orientação técnica e acadêmica, à instituição de ensino pelo suporte, a colegas que colaboraram em discussões ou na coleta de dados, além de familiares e amigos pelo apoio emocional. Também podem ser mencionados órgãos de fomento ou empresas que ofereceram bolsas ou suporte financeiro.\n"
            "A principal função dos agradecimentos é reconhecer publicamente o suporte e a colaboração recebidos, demonstrando respeito e humildade diante das contribuições que tornaram possível a elaboração do trabalho.")
            # O título será adicionado na cor preta e o texto na cor vermelha, pois ele é uma explicação e um exemplo do que é um texto de agradecimento e deve ser apagado posteriormente:
            formatacao(documento, "AGRADECIMENTOS", True, "nada", "centralizado")
            formatacao(documento, explicacaoAgradecimento, False, "nada", "esquerda", 12, 0, 0, (255, 0, 0))
        # Pula para a próxima página:
        documento.add_page_break()


    # Próxima Página: Epígrafe (Opcional):

    # Verificar se o usuário deseja que o seu documento tenha uma página contendo uma epígrafe:
    print("Epígrafe = Frase curta, citação ou trecho de obra colocado no início de um trabalho acadêmico, capítulo ou livro.\nEla serve para refletir o tema, inspirar o leitor ou expressar um sentimento do autor sobre o conteúdo que será apresentado.")
    if verificarDesejoDoUsuario("UMA PÁGINA CONTENDO UMA EPÍGRAFE") == 1:
        # Verificar em qual momento o usuário deseja adicionar o conteúdo:
        if decidirMomentoDeAdicao("a EPÍGRAFE") == 1:
            epigrafe = input("Digite a sua epígrafe:\n" 
            "Dicas:\n"
            "->Redija uma frase curta, citação ou trecho de obra que você goste, não preciasa ter sentido com o tema do trabalho;\n" 
            "->Nunca se esqueça de ter colacar o autor."
            "Exemplo:\n"
            "“Conhece-te a ti mesmo.”\n"
            "— Sócrates\n"
            "-> ")
            formatacao(documento, epigrafe, False, "nada", "justificado", 12, 22, 7.5)  
        else:
            explicacaoEpigrafe = ("A epígrafe é um elemento opcional nos trabalhos acadêmicos e consiste em uma citação, frase curta ou trecho de obra escolhido pelo autor.\n"
                "Ela costuma ser posicionada em uma página exclusiva, após a folha de rosto e antes do sumário ou do resumo, e não possui título.\n"
                "A epígrafe pode ter relação com o tema do trabalho, mas também pode apenas refletir um sentimento, pensamento ou inspiração pessoal do autor.\n"
                "É importante incluir o nome do autor da frase utilizada.\n"
                "*ESSA PÁGINA NÃO CONTÉM NADA ALÉM DO TEXTO, OU SEJA, NÃO POSSUI TÍTULOS, IMAGENS OU OUTROS ELEMENTOS GRÁFICOS.\n"
                "*A citação é geralmente posicionada no canto inferior direito da página.\n"
                "*Substítua o exmplo, pela sua epígrafe!")
            epigrafe = ("“A educação é a arma mais poderosa que você pode usar para mudar o mundo.”\n"
                "— Nelson Mandela")
            formatacao(documento, explicacaoEpigrafe, True, "nada", "esquerda", 12, 0, 0, (255, 0, 0))  
            formatacao(documento, epigrafe, False, "nada", "justificado", 12, 8, 7.5)  
        # Pula para a próxima página:
        documento.add_page_break()

    # Próxima Página: Resumo (Obrigatório):

    # ADICIONAR FUNÇÃO PARA O USUÁRIO ESCOLHER SE DESEJA ESCREVER AGORA OU MAIS TARDE!!!!!!!!!!
    if decidirMomentoDeAdicao("o RESUMO DA OBRA") == 1:
        resumo = input("Digite o resumo da obra:\n"
        "Dicas:\n"
        "-> Apresente o tema principal da obra;\n"
        "-> Mencione os personagens principais (se houver);\n"
        "-> Destaque os acontecimentos mais importantes;\n"
        "-> Fale sobre o desfecho, sem muitos spoilers se não for necessário;\n"
        "-> Seja claro e objetivo;\n"
        "-> Use suas próprias palavras para evitar plágio.\n"
        "-> ")
        formatacao(documento, "RESUMO", True, "centralizado")
        formatacao(documento, resumo, False, "nada", "justificado", 12)
    else:
        explicacaoResumo = ("O resumo é uma apresentação concisa do conteúdo de um trabalho acadêmico, elaborado segundo a norma ABNT NBR 6028:2021.\n"
        "Seu objetivo é oferecer ao leitor uma visão geral clara e rápida sobre o tema abordado, os objetivos, a metodologia, os resultados e as conclusões da pesquisa.\n\n"

        "O que deve conter no resumo:\n"
        "1. Tema do trabalho – assunto central tratado;\n"
        "2. Objetivo(s) – o que o trabalho se propôs a alcançar;\n"
        "3. Metodologia – como a pesquisa foi realizada (ex: pesquisa bibliográfica, estudo de caso etc.);\n"
        "4. Resultados – principais descobertas ou análises;\n"
        "5. Conclusão – síntese final ou implicações dos resultados.\n\n"

        "Regras importantes:\n"
        "- Deve ter até 250 palavras para trabalhos acadêmicos;\n"
        "- Deve ser escrito em um único parágrafo (sem separações);\n"
        "- Não pode conter citações diretas, imagens, gráficos, nem abreviações não explicadas;\n"
        "- Deve estar redigido em terceira pessoa, com verbos na voz ativa e no tempo passado;\n"
        "- As palavras-chave devem vir logo abaixo do resumo, de 3 a 5, separadas por ponto e vírgula, finalizando com ponto.\n\n"

        "Exemplo de frase inicial no resumo:\n"
        "\"Este trabalho analisou o impacto das redes sociais no desempenho acadêmico de estudantes do ensino médio...\"\n")
        formatacao(documento, "RESUMO", True, True, "centralizado")
        formatacao(documento, explicacaoResumo, True, "nada", "esquerda", 12, 0, 0, (255, 0, 0))

    # Pula para a próxima página
    documento.add_page_break()

    # Próxima Página: Abstract (Obrigatório):
    formatacao(documento, "ABSTRACT", True, True, "centralizado")
    formatacao(documento, "VERSÃO DO SEU RESUMO EM INGLÊS", True, "nada", "justificado", 12, 0, 0, (255, 0, 0))

    # Pula para a próxima página
    documento.add_page_break()

    # Próxima Página: Lista de Ilustrações (Opcional):

    # Verificar se o usuário deseja que o seu documento tenha uma página dedicada a lista de ilustrações:
    if verificarDesejoDoUsuario("UMA PÁGINA CONTENDO UMA LISTA DE ILUSTRAÇÕES") == 1:
        formatacao(documento, "LISTA DE ILUSTRAÇÕES", True, "nada", "centralizado")
        # Pula para a próxima página:
        documento.add_page_break()

    # Próxima Página: Lista de Tabelas (Opcional):

    # Verificar se o usuário deseja que o seu documento tenha uma página dedicada a lista de tabelas:
    if verificarDesejoDoUsuario("UMA PÁGINA CONTENDO UMA LISTA DE TABELAS") == 1:
        formatacao(documento, "LISTA DE TABELAS", True, "nada", "centralizado")
        # Pula para a próxima página:
        documento.add_page_break()

    # Próxima Página: Lista de Abreviaturas e Siglas (Opcional):

    # Verificar se o usuário deseja que o seu documento tenha uma página dedicada a lista de abreviaturas e siglass:
    if verificarDesejoDoUsuario("UMA PÁGINA CONTENDO UMA LISTA DE ABREVIATURAS E SIGLAS") == 1:
        formatacao(documento, "LISTA DE ABREVIATURAS E SIGLAS", True, "nada", "centralizado")
        # Pula para a próxima página:
        documento.add_page_break()

    # Próxima Página: Lista de Símbolos (Opcional):

    # Verificar se o usuário deseja que o seu documento tenha uma página dedicada a lista de símbolos:
    if verificarDesejoDoUsuario("UMA PÁGINA CONTENDO UMA LISTA DE SÍMBOLOS") == 1:
        formatacao(documento, "LISTA DE SÍMBOLOS", True, True, "centralizado")
        # Pula para a próxima página:
        documento.add_page_break()

    # Próxima Página: Sumário (Obrigatório):
    formatacao(documento, "SUMÁRIO", True, True, "centralizado")

    # Próxima Página: Introdução (Obrigatório):
    documento.add_heading('1 INTRODUÇÃO', level=1)

    # Próxima Página: Desenvolvimento (Obrigatório):
    documento.add_heading('2 DESENVOLVIMENTO', level=1)

    # Próxima Página: Conclusão (Obrigatório):
    documento.add_heading('3 CONCLUSÃO', level=1)

    # Próxima Página: Referências (Obrigatório)
    p_referencias = documento.add_heading('REFERÊNCIAS', level=1)
    p_referencias.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Próxima Página: Glossário (Opcional)
    if verificarDesejoDoUsuario("UMA PÁGINA CONTENDO UM GLOSSÁRIO") == 1:
        print("Ok!!! Você DESEJA adicionar um GLOSSÁRIO ao seu trabalho!!!")

        p_glossario = documento.add_heading('GLOSSÁRIO', level=1)
        p_glossario.alignment = WD_ALIGN_PARAGRAPH.CENTER


    # Próxima Página: Apêndice A (Opcional)
    if verificarDesejoDoUsuario("um APÊNDICE") == 1:
        print("Ok!!! Você DESEJA adicionar um APÊNDICE ao seu trabalho!!!")
    
        p_apendice = documento.add_heading('APÊNDICE A', level=1)
        p_apendice.alignment = WD_ALIGN_PARAGRAPH.CENTER

        formatacao(documento, "O apêndice é um elemento pós-textual que contém materiais produzidos pelo próprio autor do trabalho, com o objetivo de complementar ou detalhar informações abordadas no corpo do texto. Trata-se de conteúdos que, embora relevantes, poderiam prejudicar o fluxo da leitura caso fossem incluídos diretamente no desenvolvimento. Por isso, são apresentados separadamente ao final do documento. Entre os exemplos de apêndice, podemos citar um questionário elaborado pelo autor, um roteiro de entrevista, cálculos detalhados ou até mesmo um código-fonte desenvolvido por quem realizou o trabalho. Cada apêndice é identificado por letras maiúsculas consecutivas (Apêndice A, Apêndice B etc.), seguido do título explicativo. A formatação deve apresentar o título centralizado, em letras maiúsculas e negrito.", True, True, "esquerda", cor = (255, 0, 0))


    # Próxima Página: Anexo A (Opcional)
    if verificarDesejoDoUsuario("um ANEXO") == 1:
        print("Ok!!! Você DESEJA adicionar um ANEXO ao seu trabalho!!!")
        
        p_anexo = documento.add_heading('ANEXO A', level=1)
        p_anexo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        formatacao(documento, "O anexo, por sua vez, também é um elemento pós-textual, mas difere do apêndice pelo fato de conter materiais que não foram produzidos pelo autor do trabalho. São documentos ou conteúdos externos, reunidos com a finalidade de reforçar, ilustrar ou comprovar informações que aparecem ao longo do texto. Exemplos comuns de anexos incluem leis, trechos de livros, reportagens, documentos oficiais, prints de páginas da internet ou tabelas divulgadas por instituições reconhecidas. Assim como os apêndices, os anexos são identificados por letras maiúsculas consecutivas (Anexo A, Anexo B etc.) e devem apresentar um título claro e objetivo. A formatação também exige que o título fique centralizado, em letras maiúsculas e em negrito.", True, True, "esquerda", cor = (255, 0, 0))


    print("\n-----------------------------------------------------------------------------------------------------------\nCHEGAMOS AO FINAL! AGORA É A HORA DE NOMEARMOS O SEU ARQUIVO!\n")
    nome_do_documento = input("Escolha um nome para o seu arquivo:\n-> ")
    documento.save(nome_do_documento + ".docx")
except ImportError:
    print("⚠️ A biblioteca 'python-docx' não está instalada.")
    print("Para instalar, execute no prompt: pip install python-docx")

